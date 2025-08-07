#!/usr/bin/env python3
"""
Certificate of Completion (COC) Generator for Advanced Machine Co.
Generates COC PDFs by integrating QuickBooks Online data with local database.

Author: Advanced Machine Co. MRP System
Created: August 2025
"""

import os
import sys
import logging
from datetime import datetime, date
from typing import Dict, Optional, Tuple
import mysql.connector
from mysql.connector import Error
from docx import Document
import pypandoc
import tempfile
import shutil

# QuickBooks integration (requires intuitlib)
try:
    from intuitlib.client import AuthClient
    from intuitlib.enums import Scopes
    from quickbooks import QuickBooks
    from quickbooks.objects import Invoice, Customer, Item
    QB_AVAILABLE = True
except ImportError:
    QB_AVAILABLE = False
    print("Warning: QuickBooks libraries not installed. Install with: pip install python-quickbooks intuitlib")


class COCGenerator:
    """
    Certificate of Completion Generator
    
    Handles the complete workflow of generating COC PDFs:
    1. Fetch data from QuickBooks Online and local database
    2. Fill DOCX template with actual data
    3. Convert to PDF
    4. Log certificate to database
    """
    
    def __init__(self, config: Dict):
        """
        Initialize COC Generator with configuration
        
        Args:
            config: Dictionary containing database and QuickBooks configuration
        """
        self.config = config
        self.logger = self._setup_logging()
        self.db_connection = None
        self.qb_client = None
        
        # Initialize connections
        self._connect_database()
        if QB_AVAILABLE:
            self._connect_quickbooks()
    
    def _setup_logging(self) -> logging.Logger:
        """Setup logging configuration"""
        logger = logging.getLogger('COCGenerator')
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def _connect_database(self):
        """Establish MySQL database connection"""
        try:
            self.db_connection = mysql.connector.connect(
                host=self.config['database']['host'],
                database=self.config['database']['database'],
                user=self.config['database']['user'],
                password=self.config['database']['password'],
                port=self.config['database'].get('port', 3306)
            )
            self.logger.info("Database connection established")
        except Error as e:
            self.logger.error(f"Database connection failed: {e}")
            raise
    
    def _connect_quickbooks(self):
        """Establish QuickBooks Online connection"""
        if not QB_AVAILABLE:
            self.logger.warning("QuickBooks libraries not available")
            return
        
        try:
            # Initialize QuickBooks client
            # Note: This requires OAuth setup - see setup instructions below
            auth_client = AuthClient(
                client_id=self.config['quickbooks']['client_id'],
                client_secret=self.config['quickbooks']['client_secret'],
                environment=self.config['quickbooks'].get('environment', 'sandbox'),
                redirect_uri=self.config['quickbooks']['redirect_uri']
            )
            
            self.qb_client = QuickBooks(
                auth_client=auth_client,
                refresh_token=self.config['quickbooks']['refresh_token'],
                company_id=self.config['quickbooks']['company_id']
            )
            self.logger.info("QuickBooks connection established")
        except Exception as e:
            self.logger.error(f"QuickBooks connection failed: {e}")
            # Continue without QB - can still generate COCs with manual data
    
    def get_work_order_data(self, work_order_id: int) -> Dict:
        """
        Fetch work order data from database
        
        Args:
            work_order_id: Work order ID
            
        Returns:
            Dictionary containing work order details
        """
        try:
            cursor = self.db_connection.cursor(dictionary=True)
            
            # Query work order with customer and part details
            query = """
            SELECT 
                wo.WorkOrderID,
                wo.WorkOrderNumber,
                wo.CustomerPONumber,
                wo.QuantityOrdered,
                wo.QuantityCompleted,
                c.CustomerName,
                c.QuickBooksID as CustomerQBID,
                p.PartNumber,
                p.PartName,
                p.Description,
                p.DrawingNumber,
                p.FSN
            FROM WorkOrders wo
            JOIN Customers c ON wo.CustomerID = c.CustomerID
            JOIN Parts p ON wo.PartID = p.PartID
            WHERE wo.WorkOrderID = %s
            """
            
            cursor.execute(query, (work_order_id,))
            result = cursor.fetchone()
            
            if not result:
                raise ValueError(f"Work order {work_order_id} not found")
            
            # Get final quantity from production stages (latest "ready to ship" stage)
            quantity_query = """
            SELECT QuantityOut 
            FROM ProductionStages 
            WHERE WorkOrderID = %s 
            AND StageName LIKE '%ready to ship%' 
            ORDER BY StageDate DESC 
            LIMIT 1
            """
            
            cursor.execute(quantity_query, (work_order_id,))
            quantity_result = cursor.fetchone()
            
            # Use production stage quantity if available, otherwise use completed quantity
            final_quantity = (quantity_result['QuantityOut'] if quantity_result 
                            else result['QuantityCompleted'])
            
            result['FinalQuantity'] = final_quantity
            
            cursor.close()
            return result
            
        except Error as e:
            self.logger.error(f"Database query failed: {e}")
            raise
    
    def get_quickbooks_invoice_data(self, invoice_id: str) -> Dict:
        """
        Fetch invoice data from QuickBooks Online
        
        Args:
            invoice_id: QuickBooks invoice ID
            
        Returns:
            Dictionary containing invoice details
        """
        if not self.qb_client:
            self.logger.warning("QuickBooks not available - using database data only")
            return {}
        
        try:
            # Fetch invoice from QuickBooks
            invoice = Invoice.get(invoice_id, qb=self.qb_client)
            
            # Extract relevant data
            invoice_data = {
                'InvoiceNumber': invoice.DocNumber,
                'InvoiceDate': invoice.TxnDate,
                'CustomerRef': invoice.CustomerRef,
                'LineItems': []
            }
            
            # Process line items
            for line in invoice.Line:
                if hasattr(line, 'SalesItemLineDetail'):
                    item_data = {
                        'ItemRef': line.SalesItemLineDetail.ItemRef,
                        'Qty': line.SalesItemLineDetail.Qty,
                        'Description': line.Description
                    }
                    invoice_data['LineItems'].append(item_data)
            
            return invoice_data
            
        except Exception as e:
            self.logger.error(f"QuickBooks query failed: {e}")
            return {}
    
    def fill_template(self, template_path: str, data: Dict) -> str:
        """
        Fill DOCX template with actual data
        
        Args:
            template_path: Path to DOCX template file
            data: Dictionary containing data to fill
            
        Returns:
            Path to filled DOCX file
        """
        try:
            # Load template
            doc = Document(template_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # Save filled document
            temp_dir = tempfile.mkdtemp()
            filled_docx_path = os.path.join(temp_dir, f"COC_filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(filled_docx_path)
            
            return filled_docx_path
            
        except Exception as e:
            self.logger.error(f"Template filling failed: {e}")
            raise
    
    def convert_to_pdf(self, docx_path: str, output_dir: str) -> str:
        """
        Convert DOCX to PDF using pypandoc
        
        Args:
            docx_path: Path to DOCX file
            output_dir: Directory to save PDF
            
        Returns:
            Path to generated PDF file
        """
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate PDF filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_filename = f"COC_{timestamp}.pdf"
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            # Convert DOCX to PDF using pypandoc
            pypandoc.convert_file(
                docx_path, 
                'pdf', 
                outputfile=pdf_path,
                extra_args=['--pdf-engine=xelatex']  # Use xelatex for better formatting
            )
            
            self.logger.info(f"PDF generated: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            self.logger.error(f"PDF conversion failed: {e}")
            # Try without extra args if xelatex fails
            try:
                self.logger.info("Retrying PDF conversion without xelatex...")
                pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
                self.logger.info(f"PDF generated (fallback): {pdf_path}")
                return pdf_path
            except Exception as e2:
                self.logger.error(f"PDF conversion fallback also failed: {e2}")
                raise
    
    def log_certificate(self, work_order_data: Dict, pdf_path: str, created_by: str = "System") -> int:
        """
        Log certificate details to database
        
        Args:
            work_order_data: Work order data dictionary
            pdf_path: Path to generated PDF
            created_by: User who created the certificate
            
        Returns:
            Certificate log ID
        """
        try:
            cursor = self.db_connection.cursor()
            
            # Generate certificate number
            cert_number = f"COC-{work_order_data['WorkOrderNumber']}-{datetime.now().strftime('%Y%m%d')}"
            
            # Insert into CertificatesLog
            insert_query = """
            INSERT INTO CertificatesLog (
                CertificateNumber, WorkOrderID, CustomerID, PartNumber, 
                Description, CustomerPONumber, Quantity, CompletionDate,
                DrawingNumber, FSN, DocumentPath, CreatedBy
            ) VALUES (
                %s, %s, (SELECT CustomerID FROM WorkOrders WHERE WorkOrderID = %s),
                %s, %s, %s, %s, %s, %s, %s, %s, %s
            )
            """
            
            values = (
                cert_number,
                work_order_data['WorkOrderID'],
                work_order_data['WorkOrderID'],  # For subquery
                work_order_data['PartNumber'],
                work_order_data['Description'] or work_order_data['PartName'],
                work_order_data['CustomerPONumber'],
                work_order_data['FinalQuantity'],
                date.today(),
                work_order_data['DrawingNumber'],
                work_order_data['FSN'],
                pdf_path,
                created_by
            )
            
            cursor.execute(insert_query, values)
            self.db_connection.commit()
            
            cert_log_id = cursor.lastrowid
            cursor.close()
            
            self.logger.info(f"Certificate logged with ID: {cert_log_id}")
            return cert_log_id
            
        except Error as e:
            self.logger.error(f"Certificate logging failed: {e}")
            self.db_connection.rollback()
            raise
    
    def generate_coc(self, work_order_id: int, invoice_id: Optional[str] = None, 
                     created_by: str = "System") -> Tuple[str, int]:
        """
        Main method to generate Certificate of Completion
        
        Args:
            work_order_id: Work order ID
            invoice_id: Optional QuickBooks invoice ID
            created_by: User generating the certificate
            
        Returns:
            Tuple of (PDF file path, certificate log ID)
        """
        try:
            self.logger.info(f"Generating COC for Work Order {work_order_id}")
            
            # Get work order data
            work_order_data = self.get_work_order_data(work_order_id)
            
            # Get QuickBooks data if available
            qb_data = {}
            if invoice_id and self.qb_client:
                qb_data = self.get_quickbooks_invoice_data(invoice_id)
            
            # Prepare template data
            template_data = {
                'DATE': date.today().strftime('%m-%d-%Y'),
                'DESCRIPTION': work_order_data['Description'] or work_order_data['PartName'],
                'IAW_SPEC_DWG': work_order_data['PartNumber'],
                'QUANTITY': work_order_data['FinalQuantity'],
                'PO': work_order_data['CustomerPONumber']
            }
            
            # Fill template
            template_path = self.config['template']['path']
            filled_docx_path = self.fill_template(template_path, template_data)
            
            # Convert to PDF
            output_dir = self.config['output']['directory']
            pdf_path = self.convert_to_pdf(filled_docx_path, output_dir)
            
            # Log certificate
            cert_log_id = self.log_certificate(work_order_data, pdf_path, created_by)
            
            # Cleanup temporary files
            if os.path.exists(filled_docx_path):
                os.remove(filled_docx_path)
            
            self.logger.info(f"COC generation completed: {pdf_path}")
            return pdf_path, cert_log_id
            
        except Exception as e:
            self.logger.error(f"COC generation failed: {e}")
            raise
    
    def close_connections(self):
        """Close database and QuickBooks connections"""
        if self.db_connection and self.db_connection.is_connected():
            self.db_connection.close()
            self.logger.info("Database connection closed")


def load_config() -> Dict:
    """
    Load configuration from environment variables or config file
    
    Returns:
        Configuration dictionary
    """
    return {
        'database': {
            'host': os.getenv('DB_HOST', 'localhost'),
            'database': os.getenv('DB_NAME', 'amcmrp'),
            'user': os.getenv('DB_USER', 'amc'),
            'password': os.getenv('DB_PASSWORD', 'Workbench.lavender.chrome'),
            'port': int(os.getenv('DB_PORT', 3306))
        },
        'quickbooks': {
            'client_id': os.getenv('QB_CLIENT_ID', ''),
            'client_secret': os.getenv('QB_CLIENT_SECRET', ''),
            'redirect_uri': os.getenv('QB_REDIRECT_URI', 'https://developer.intuit.com/v2/OAuth2Playground/RedirectUrl'),
            'refresh_token': os.getenv('QB_REFRESH_TOKEN', ''),
            'company_id': os.getenv('QB_COMPANY_ID', ''),
            'environment': os.getenv('QB_ENVIRONMENT', 'sandbox')  # or 'production'
        },
        'template': {
            'path': os.getenv('COC_TEMPLATE_PATH', '../Dev Assets/COC Template.docx')
        },
        'output': {
            'directory': os.getenv('COC_OUTPUT_DIR', '../generated_cocs')
        }
    }


def main():
    """
    Example usage of COC Generator
    """
    try:
        # Load configuration
        config = load_config()
        
        # Initialize generator
        coc_gen = COCGenerator(config)
        
        # Example: Generate COC for work order ID 1
        work_order_id = 1
        pdf_path, cert_id = coc_gen.generate_coc(work_order_id, created_by="Test User")
        
        print(f"COC generated successfully!")
        print(f"PDF Path: {pdf_path}")
        print(f"Certificate ID: {cert_id}")
        
        # Close connections
        coc_gen.close_connections()
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()


"""
SETUP INSTRUCTIONS:

1. Install Required Dependencies:
   pip install mysql-connector-python python-docx pypandoc python-quickbooks intuitlib
   
   Note: pypandoc requires pandoc to be installed on your system:
   - macOS: brew install pandoc
   - Ubuntu/Debian: sudo apt-get install pandoc
   - Windows: Download from https://pandoc.org/installing.html

2. Environment Variables:
   Set the following environment variables or modify the load_config() function:
   
   Database:
   - DB_HOST: MySQL host (default: localhost)
   - DB_NAME: Database name (default: amc_mrp_system)
   - DB_USER: Database user (default: root)
   - DB_PASSWORD: Database password
   - DB_PORT: Database port (default: 3306)
   
   QuickBooks (optional):
   - QB_CLIENT_ID: QuickBooks app client ID
   - QB_CLIENT_SECRET: QuickBooks app client secret
   - QB_REDIRECT_URI: OAuth redirect URI
   - QB_REFRESH_TOKEN: OAuth refresh token
   - QB_COMPANY_ID: QuickBooks company ID
   - QB_ENVIRONMENT: 'sandbox' or 'production'
   
   Paths:
   - COC_TEMPLATE_PATH: Path to DOCX template file
   - COC_OUTPUT_DIR: Directory to save generated PDFs

3. QuickBooks Setup:
   - Create a QuickBooks Online app at https://developer.intuit.com/
   - Get OAuth tokens using the OAuth playground
   - Store tokens securely

4. Template Setup:
   - Ensure the DOCX template exists with placeholders: {{DATE}}, {{DESCRIPTION}}, {{IAW_SPEC_DWG}}, {{QUANTITY}}, {{PO}}

5. Database Setup:
   - Ensure the database schema is created (use amc_mrp_schema.sql)
   - Verify database connectivity

USAGE EXAMPLES:

# Basic usage
from cocGenerate import COCGenerator, load_config

config = load_config()
coc_gen = COCGenerator(config)

# Generate COC for work order
pdf_path, cert_id = coc_gen.generate_coc(work_order_id=123)

# Generate COC with QuickBooks invoice
pdf_path, cert_id = coc_gen.generate_coc(work_order_id=123, invoice_id="QB_INV_456")

# Close connections
coc_gen.close_connections()
"""
