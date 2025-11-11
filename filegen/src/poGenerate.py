#!/usr/bin/env python3
"""
Internal Purchase Order (PO) Generator for Advanced Machine Co.
Generates PO PDFs using local database data.

Author: Advanced Machine Co. MRP System
Created: August 2025
"""

import os
import sys
import logging
from datetime import datetime, date
from typing import Dict, Optional, Tuple, List
import mysql.connector
from mysql.connector import Error
from docx import Document
import pypandoc
import tempfile
import shutil

# No QuickBooks integration - using database only

class POGenerator:
    """
    Internal Purchase Order Generator
    
    Handles the complete workflow of generating Internal PO PDFs:
    1. Fetch BOM process data from local database
    2. Get vendor data from local database
    3. Fill DOCX template with actual data
    4. Convert to PDF
    5. Log PO to database
    """
    
    def __init__(self, config: Dict):
        """
        Initialize PO Generator with configuration
        
        Args:
            config: Dictionary containing database configuration
        """
        self.config = config
        self.logger = self._setup_logging()
        self.db_connection = None
        
        # Initialize database connection
        self._connect_database()
    
    def _setup_logging(self) -> logging.Logger:
        """Setup logging configuration"""
        logger = logging.getLogger('POGenerator')
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def _connect_database(self):
        """Establish MySQL database connection"""
        try:
            self.db_connection = mysql.connector.connect(
                host=self.config['database']['host'],
                database=self.config['database']['database'],
                user=self.config['database']['user'],
                password=self.config['database']['password'],
                port=self.config['database'].get('port', 3306)
            )
            self.logger.info("Database connection established")
        except Error as e:
            self.logger.error(f"Database connection failed: {e}")
            raise
    
    
    def get_bom_process_data(self, process_id: int) -> Dict:
        """
        Fetch BOM process data with related work order and part information
        
        Args:
            process_id: BOM Process ID
            
        Returns:
            Dictionary containing process, work order, and part details
        """
        try:
            cursor = self.db_connection.cursor(dictionary=True)
            
            # Query BOM process with work order, part, and vendor details
            query = """
            SELECT 
                bp.ProcessID,
                bp.ProcessType,
                bp.ProcessName,
                bp.Quantity,
                bp.UnitOfMeasure,
                bp.EstimatedCost,
                bp.ActualCost,
                bp.LeadTimeDays,
                bp.CertificationRequired,
                bp.ProcessRequirements,
                bp.Status as ProcessStatus,
                wo.WorkOrderID,
                wo.CustomerPONumber,
                p.PartNumber,
                p.Description,
                p.Material,
                v.VendorID,
                v.VendorName,
                v.QuickBooksID as VendorQBID,
                v.ContactPhone,
                v.ContactEmail,
                v.Address as VendorAddress,
                c.CustomerName
            FROM BOMProcesses bp
            JOIN BOM b ON bp.BOMID = b.BOMID
            JOIN WorkOrders wo ON b.WorkOrderID = wo.WorkOrderID
            JOIN Parts p ON wo.PartID = p.PartID
            JOIN Customers c ON wo.CustomerID = c.CustomerID
            LEFT JOIN Vendors v ON bp.VendorID = v.VendorID
            WHERE bp.ProcessID = %s
            """
            
            cursor.execute(query, (process_id,))
            result = cursor.fetchone()
            
            if not result:
                raise ValueError(f"BOM Process {process_id} not found")
            
            cursor.close()
            return result
            
        except Error as e:
            self.logger.error(f"Database query failed: {e}")
            raise
    
    
    def generate_po_number(self) -> str:
        """
        Generate sequential PO number in format MMDDYY-XX
        
        Returns:
            Generated PO number
        """
        try:
            cursor = self.db_connection.cursor()
            
            # Get today's date in MMDDYY format
            today = date.today()
            date_prefix = today.strftime('%m%d%y')
            
            # Find the highest sequence number for today
            query = """
            SELECT PONumber FROM PurchaseOrdersLog 
            WHERE PONumber LIKE %s 
            ORDER BY PONumber DESC 
            LIMIT 1
            """
            
            cursor.execute(query, (f"{date_prefix}-%",))
            result = cursor.fetchone()
            
            if result:
                # Extract sequence number and increment
                last_po = result[0]
                sequence = int(last_po.split('-')[1]) + 1
            else:
                # First PO of the day
                sequence = 1
            
            po_number = f"{date_prefix}-{sequence:02d}"
            cursor.close()
            
            return po_number
            
        except Error as e:
            self.logger.error(f"PO number generation failed: {e}")
            # Fallback to timestamp-based number
            return datetime.now().strftime('%m%d%y-%H%M')
    
    def fill_po_template(self, template_path: str, data: Dict) -> str:
        """
        Fill PO DOCX template with actual data while preserving formatting
        
        Args:
            template_path: Path to DOCX template file
            data: Dictionary containing data to fill
            
        Returns:
            Path to filled DOCX file
        """
        try:
            # Load template
            doc = Document(template_path)
            
            # Replace placeholders in paragraphs while preserving formatting
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        # Find runs that contain the placeholder
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                # Replace while preserving the run's formatting
                                run.text = run.text.replace(placeholder, str(value))
                        
                        # Handle multi-run placeholders
                        if placeholder in paragraph.text and not any(placeholder in run.text for run in paragraph.runs):
                            full_text = paragraph.text
                            if placeholder in full_text:
                                new_text = full_text.replace(placeholder, str(value))
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size
                                    bold = first_run.font.bold
                                    italic = first_run.font.italic
                                    
                                    for run in paragraph.runs[::-1]:
                                        paragraph._element.remove(run._element)
                                    
                                    new_run = paragraph.add_run(new_text)
                                    if font_name:
                                        new_run.font.name = font_name
                                    if font_size:
                                        new_run.font.size = font_size
                                    if bold:
                                        new_run.font.bold = bold
                                    if italic:
                                        new_run.font.italic = italic
            
            # Replace placeholders in tables while preserving formatting
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
                                    
                                    if placeholder in paragraph.text and not any(placeholder in run.text for run in paragraph.runs):
                                        full_text = paragraph.text
                                        if placeholder in full_text:
                                            new_text = full_text.replace(placeholder, str(value))
                                            if paragraph.runs:
                                                first_run = paragraph.runs[0]
                                                font_name = first_run.font.name
                                                font_size = first_run.font.size
                                                bold = first_run.font.bold
                                                italic = first_run.font.italic
                                                
                                                for run in paragraph.runs[::-1]:
                                                    paragraph._element.remove(run._element)
                                                
                                                new_run = paragraph.add_run(new_text)
                                                if font_name:
                                                    new_run.font.name = font_name
                                                if font_size:
                                                    new_run.font.size = font_size
                                                if bold:
                                                    new_run.font.bold = bold
                                                if italic:
                                                    new_run.font.italic = italic
            
            # Save filled document
            temp_dir = tempfile.mkdtemp()
            filled_docx_path = os.path.join(temp_dir, f"PO_filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(filled_docx_path)
            
            return filled_docx_path
            
        except Exception as e:
            self.logger.error(f"Template filling failed: {e}")
            raise
    
    def convert_to_pdf(self, docx_path: str, output_dir: str) -> str:
        """
        Convert DOCX to PDF using LibreOffice with enhanced font preservation
        
        Args:
            docx_path: Path to DOCX file
            output_dir: Directory to save PDF
            
        Returns:
            Path to generated PDF file
        """
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate PDF filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_filename = f"PO_{timestamp}.pdf"
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            # Use LibreOffice to convert DOCX to PDF (headless mode)
            import subprocess
            
            libreoffice_cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                '-env:UserInstallation=file:///tmp/libreoffice_profile',
                docx_path
            ]
            
            self.logger.info(f"Converting DOCX to PDF using LibreOffice: {docx_path}")
            result = subprocess.run(libreoffice_cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
            # LibreOffice creates PDF with same name as input file but .pdf extension
            docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
            libreoffice_pdf = os.path.join(output_dir, f"{docx_basename}.pdf")
            
            # Rename to our desired filename
            if os.path.exists(libreoffice_pdf):
                if libreoffice_pdf != pdf_path:
                    shutil.move(libreoffice_pdf, pdf_path)
                self.logger.info(f"PDF generated successfully: {pdf_path}")
                return pdf_path
            else:
                raise Exception(f"LibreOffice PDF output not found: {libreoffice_pdf}")
            
        except Exception as e:
            self.logger.error(f"LibreOffice PDF conversion failed: {e}")
            # Fallback to pypandoc if LibreOffice fails
            try:
                self.logger.info("Falling back to pypandoc conversion...")
                pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
                self.logger.info(f"PDF generated (pypandoc fallback): {pdf_path}")
                return pdf_path
            except Exception as e2:
                self.logger.error(f"Pypandoc fallback also failed: {e2}")
                raise Exception(f"Both LibreOffice and pypandoc conversion failed. LibreOffice: {e}, Pypandoc: {e2}")
    
    
    def log_purchase_order(self, process_data: Dict, po_number: str, pdf_path: str, 
                          qb_po_id: Optional[str] = None, created_by: str = "System") -> int:
        """
        Log purchase order details to database
        
        Args:
            process_data: BOM process data dictionary
            po_number: Generated PO number
            pdf_path: Path to generated PDF
            qb_po_id: QuickBooks PO ID if created
            created_by: User who created the PO
            
        Returns:
            PO log ID
        """
        try:
            cursor = self.db_connection.cursor()
            
            # Insert into PurchaseOrdersLog
            insert_query = """
            INSERT INTO PurchaseOrdersLog (
                PONumber, WorkOrderID, ProcessID, VendorID, PODate,
                PartNumber, Description, Material, Quantity, UnitPrice, TotalAmount,
                CertificationRequired, ProcessRequirements, Status, DocumentPath,
                CreatedBy
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
            )
            """
            
            values = (
                po_number,
                process_data['WorkOrderID'],
                process_data['ProcessID'],
                process_data['VendorID'],
                date.today(),
                process_data['PartNumber'],
                process_data['Description'],
                process_data['Material'],
                process_data['Quantity'],
                process_data['EstimatedCost'],
                process_data['EstimatedCost'] * process_data['Quantity'] if process_data['EstimatedCost'] else 0,
                process_data['CertificationRequired'],
                process_data['ProcessRequirements'],
                'Created',
                pdf_path,
                created_by
            )
            
            cursor.execute(insert_query, values)
            self.db_connection.commit()
            
            po_log_id = cursor.lastrowid
            
            # Update BOM process status to 'Ordered'
            update_query = """
            UPDATE BOMProcesses 
            SET Status = 'Ordered', UpdatedDate = CURRENT_TIMESTAMP 
            WHERE ProcessID = %s
            """
            cursor.execute(update_query, (process_data['ProcessID'],))
            self.db_connection.commit()
            
            cursor.close()
            
            self.logger.info(f"Purchase order logged with ID: {po_log_id}")
            return po_log_id
            
        except Error as e:
            self.logger.error(f"PO logging failed: {e}")
            self.db_connection.rollback()
            raise
    
    def cleanup_old_pdfs(self, output_dir: str, keep_latest: int = 5):
        """
        Clean up old PDF files, keeping only the most recent ones
        
        Args:
            output_dir: Directory containing PDF files
            keep_latest: Number of latest PDFs to keep (default: 5)
        """
        try:
            import glob
            
            # Find all PO PDF files
            pdf_pattern = os.path.join(output_dir, "PO_*.pdf")
            pdf_files = glob.glob(pdf_pattern)
            
            if len(pdf_files) > keep_latest:
                # Sort by modification time (newest first)
                pdf_files.sort(key=os.path.getmtime, reverse=True)
                
                # Remove older files
                files_to_remove = pdf_files[keep_latest:]
                for file_path in files_to_remove:
                    try:
                        os.remove(file_path)
                        self.logger.info(f"Removed old PDF: {file_path}")
                    except Exception as e:
                        self.logger.warning(f"Failed to remove old PDF {file_path}: {e}")
                        
        except Exception as e:
            self.logger.warning(f"PDF cleanup failed: {e}")
    
    def generate_internal_po(self, process_id: int, created_by: str = "System") -> Tuple[str, int]:
        """
        Main method to generate Internal Purchase Order
        
        Args:
            process_id: BOM Process ID
            created_by: User generating the PO
            
        Returns:
            Tuple of (PDF file path, PO log ID)
        """
        try:
            self.logger.info(f"Generating Internal PO for Process {process_id}")
            
            # Get BOM process data
            process_data = self.get_bom_process_data(process_id)
            
            # Generate PO number
            po_number = self.generate_po_number()
            
            # Prepare template data using database vendor information
            vendor_address = process_data.get('VendorAddress', '')
            vendor_phone = process_data.get('ContactPhone', '')
            
            # Calculate total
            unit_price = process_data.get('EstimatedCost', 0) or 0
            quantity = process_data.get('Quantity', 1)
            total_amount = unit_price * quantity
            
            template_data = {
                'PO_NUM': po_number,
                'DATE': date.today().strftime('%m/%d/%Y'),
                'VENDOR_NAME': process_data.get('VendorName', ''),
                'VENDOR_STREET': vendor_address,
                'VENDOR_ZIP': '',  # Already included in address
                'VENDOR_PHONE': vendor_phone,
                'INSTRUCTIONS': process_data.get('ProcessRequirements', ''),
                'TOTAL': f"${total_amount:.2f}" if total_amount > 0 else "TBD"
            }
            
            # Fill template
            template_path = self.config['template']['path']
            filled_docx_path = self.fill_po_template(template_path, template_data)
            
            # Convert to PDF
            output_dir = self.config['output']['directory']
            pdf_path = self.convert_to_pdf(filled_docx_path, output_dir)
            
            # Log purchase order (no QuickBooks integration)
            po_log_id = self.log_purchase_order(process_data, po_number, pdf_path, None, created_by)
            
            # Cleanup temporary files
            if os.path.exists(filled_docx_path):
                os.remove(filled_docx_path)
            
            # Clean up old PDFs (keep latest 5)
            self.cleanup_old_pdfs(output_dir, keep_latest=5)
            
            self.logger.info(f"Internal PO generation completed: {pdf_path}")
            return pdf_path, po_log_id
            
        except Exception as e:
            self.logger.error(f"Internal PO generation failed: {e}")
            raise
    
    def close_connections(self):
        """Close database and QuickBooks connections"""
        if self.db_connection and self.db_connection.is_connected():
            self.db_connection.close()
            self.logger.info("Database connection closed")


def load_config() -> Dict:
    """
    Load configuration from environment variables or config file
    
    Returns:
        Configuration dictionary
    """
    return {
        'database': {
            'host': os.getenv('DB_HOST', 'localhost'),
            'database': os.getenv('DB_NAME', 'amcmrp'),
            'user': os.getenv('DB_USER', 'amc'),
            'password': os.getenv('DB_PASSWORD', 'Workbench.lavender.chrome'),
            'port': int(os.getenv('DB_PORT', 3306))
        },
        'quickbooks': {
            'client_id': os.getenv('QB_CLIENT_ID', ''),
            'client_secret': os.getenv('QB_CLIENT_SECRET', ''),
            'redirect_uri': os.getenv('QB_REDIRECT_URI', 'https://developer.intuit.com/v2/OAuth2Playground/RedirectUrl'),
            'refresh_token': os.getenv('QB_REFRESH_TOKEN', ''),
            'company_id': os.getenv('QB_COMPANY_ID', ''),
            'environment': os.getenv('QB_ENVIRONMENT', 'sandbox'),  # or 'production'
            'create_pos': os.getenv('QB_CREATE_POS', 'false').lower() == 'true'
        },
        'template': {
            'path': os.getenv('PO_TEMPLATE_PATH', '../DevAssets/PO Template.docx')
        },
        'output': {
            'directory': os.getenv('PO_OUTPUT_DIR', './CACHE')
        }
    }


def main():
    """
    Example usage of PO Generator
    """
    try:
        # Load configuration
        config = load_config()
        
        # Initialize generator
        po_gen = POGenerator(config)
        
        # Example: Generate PO for BOM process ID 1
        process_id = 1
        pdf_path, po_id = po_gen.generate_internal_po(process_id, created_by="Test User")
        
        print(f"Internal PO generated successfully!")
        print(f"PDF Path: {pdf_path}")
        print(f"PO Log ID: {po_id}")
        
        # Close connections
        po_gen.close_connections()
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()


"""
SETUP INSTRUCTIONS:

1. Install Required Dependencies:
   pip install mysql-connector-python python-docx pypandoc python-quickbooks intuitlib
   
   Note: pypandoc requires pandoc to be installed on your system:
   - macOS: brew install pandoc
   - Ubuntu/Debian: sudo apt-get install pandoc
   - Windows: Download from https://pandoc.org/installing.html

2. Environment Variables:
   Set the following environment variables or modify the load_config() function:
   
   Database:
   - DB_HOST: MySQL host (default: localhost)
   - DB_NAME: Database name (default: amcmrp)
   - DB_USER: Database user (default: amc)
   - DB_PASSWORD: Database password
   - DB_PORT: Database port (default: 3306)
   
   QuickBooks (optional):
   - QB_CLIENT_ID: QuickBooks app client ID
   - QB_CLIENT_SECRET: QuickBooks app client secret
   - QB_REDIRECT_URI: OAuth redirect URI
   - QB_REFRESH_TOKEN: OAuth refresh token
   - QB_COMPANY_ID: QuickBooks company ID
   - QB_ENVIRONMENT: 'sandbox' or 'production'
   - QB_CREATE_POS: 'true' to create POs in QuickBooks (default: false)
   
   Paths:
   - PO_TEMPLATE_PATH: Path to DOCX template file
   - PO_OUTPUT_DIR: Directory to save generated PDFs

3. QuickBooks Setup:
   - Create a QuickBooks Online app at https://developer.intuit.com/
   - Get OAuth tokens using the OAuth playground
   - Store tokens securely
   - Ensure vendor permissions are enabled

4. Template Setup:
   - Ensure the DOCX template exists with placeholders: {{PO_NUM}}, {{DATE}}, {{VENDOR_NAME}}, etc.

5. Database Setup:
   - Ensure the database schema is created with BOMProcesses and PurchaseOrdersLog tables
   - Verify database connectivity

USAGE EXAMPLES:

# Basic usage
from poGenerate import POGenerator, load_config

config = load_config()
po_gen = POGenerator(config)

# Generate PO for BOM process
pdf_path, po_id = po_gen.generate_internal_po(process_id=123)

# Generate PO with custom user
pdf_path, po_id = po_gen.generate_internal_po(process_id=123, created_by="John Doe")

# Close connections
po_gen.close_connections()
"""
